from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"C:\Users\马肃\Documents\edunet\中越合作备忘录_中越英三语打印版.docx"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
GRAY = "666666"


def set_font(run, size=10.5, bold=False, color="000000", italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_font(run, 8.5, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页 / Trang / Page")
    set_font(run, 8.5, color=GRAY)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("谅解备忘录")
    set_font(r, 20, True, BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("BIÊN BẢN GHI NHỚ")
    set_font(r, 16, True, BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("MEMORANDUM OF UNDERSTANDING")
    set_font(r, 14, True, BLUE)


def add_intro_table(doc):
    data = [
        ("签署日期 / Ngày ký / Date", "2026年____月____日 / Ngày ____ tháng ____ năm 2026 / ____ 2026"),
        ("签署地点 / Địa điểm / Place", "____________________________________________"),
        ("甲方（越方）/ Bên A / Party A", "____________________________________________"),
        ("地址 / Địa chỉ / Address", "____________________________________________"),
        ("授权代表 / Đại diện được ủy quyền / Authorized Representative", "____________________________________________"),
        ("乙方（中方咨询公司）/ Bên B / Party B", "新瑞拓海有限公司 / Shinered Oversea Co., Ltd."),
        ("地址 / Địa chỉ / Address", "____________________________________________"),
        ("授权代表 / Đại diện được ủy quyền / Authorized Representative", "____________________________________________"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in data:
        cells = table.add_row().cells
        shade(cells[0], LIGHT_BLUE)
        for text, cell, bold in ((label, cells[0], True), (value, cells[1], False)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_font(r, 9.5, bold)
    table_geometry(table, [3000, 6360])
    doc.add_paragraph()


def add_heading(doc, zh, vi, en):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.shading = None
    r = p.add_run(f"{zh}  |  {vi}  |  {en}")
    set_font(r, 12, True, BLUE)


def add_trilingual_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["中文", "TIẾNG VIỆT", "ENGLISH"]
    for i, text in enumerate(headers):
        shade(table.rows[0].cells[i], BLUE)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, 9, True, "FFFFFF")
    set_repeat_table_header(table.rows[0])
    for zh, vi, en in rows:
        cells = table.add_row().cells
        for i, text in enumerate((zh, vi, en)):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(text)
            set_font(r, 9.2)
    table_geometry(table, [3120, 3120, 3120])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


sections = [
    ("序言", "Lời mở đầu", "Preamble", [
        ("本备忘录由上述双方于上述日期及地点签订。甲方与乙方以下单独称为“一方”，合称为“双方”。",
         "Biên bản ghi nhớ này được ký kết giữa các bên nêu trên vào ngày và tại địa điểm nêu trên. Bên A và Bên B sau đây được gọi riêng là “một Bên” và gọi chung là “các Bên”.",
         "This Memorandum of Understanding is entered into by the above parties on the date and at the place stated above. Party A and Party B are each a “Party” and together the “Parties”." ),
        ("鉴于，越中两国近年来在“长期稳定、面向未来、睦邻友好、全面合作”方针指引下，双边经贸与科技合作关系不断深化，特别是在数字经济、人工智能、绿色发展、教育科技等新兴领域，已成为两国高层共识和未来合作的重点。",
         "Xét rằng, trong những năm gần đây, dưới định hướng “ổn định lâu dài, hướng tới tương lai, láng giềng hữu nghị và hợp tác toàn diện”, quan hệ hợp tác kinh tế, thương mại và khoa học công nghệ giữa Việt Nam và Trung Quốc ngày càng sâu rộng; đặc biệt, kinh tế số, trí tuệ nhân tạo, phát triển xanh và công nghệ giáo dục đã trở thành các lĩnh vực đồng thuận cấp cao và trọng tâm hợp tác trong tương lai.",
         "Whereas, in recent years, guided by the principles of long-term stability, future orientation, good-neighborly friendship and comprehensive cooperation, economic, trade, scientific and technological cooperation between Vietnam and China has continued to deepen, particularly in emerging fields such as the digital economy, artificial intelligence, green development and education technology, which have become areas of high-level consensus and priorities for future cooperation."),
        ("鉴于，甲方作为越南国家在发展规划与高新科技研究领域的核心智库与政策推动机构，致力于引进先进技术和管理经验，服务国家数字化转型与绿色增长战略。",
         "Xét rằng, Bên A, với vai trò là tổ chức tư vấn chủ chốt và thúc đẩy chính sách của Việt Nam trong lĩnh vực quy hoạch phát triển và nghiên cứu công nghệ cao, cam kết tiếp nhận công nghệ tiên tiến và kinh nghiệm quản lý nhằm phục vụ chiến lược chuyển đổi số và tăng trưởng xanh quốc gia.",
         "Whereas, Party A, as a key Vietnamese think tank and policy-promoting institution in development planning and high-technology research, is committed to introducing advanced technologies and management experience in support of Vietnam’s national digital transformation and green growth strategies."),
        ("鉴于，乙方作为中国一家专业的商务与科技咨询公司，拥有连接中国优质科技企业资源、提供跨境法律、知识产权与金融咨询服务的专业能力。",
         "Xét rằng, Bên B là một công ty tư vấn kinh doanh và công nghệ chuyên nghiệp của Trung Quốc, có năng lực kết nối nguồn lực của các doanh nghiệp công nghệ chất lượng cao của Trung Quốc và cung cấp dịch vụ tư vấn pháp lý xuyên biên giới, sở hữu trí tuệ và tài chính.",
         "Whereas, Party B is a professional business and technology consulting company of China, with expertise in connecting high-quality Chinese technology enterprises and providing cross-border legal, intellectual property and financial advisory services."),
        ("鉴于，双方有意建立一种非官方、非约束性的合作框架，以促进信息交流、资源共享并最终推动实质性的企业项目合作。为此，双方经友好协商，达成以下谅解：",
         "Xét rằng, các Bên mong muốn thiết lập một khuôn khổ hợp tác phi chính thức và không ràng buộc nhằm thúc đẩy trao đổi thông tin, chia sẻ nguồn lực và cuối cùng thúc đẩy hợp tác dự án doanh nghiệp thực chất. Trên cơ sở tham vấn hữu nghị, các Bên thống nhất như sau:",
         "Whereas, the Parties intend to establish an unofficial and non-binding cooperation framework to promote information exchange, resource sharing and, ultimately, substantive enterprise project cooperation. Accordingly, through friendly consultation, the Parties reach the following understanding:")
    ]),
    ("第一条：合作目标", "Điều 1: Mục tiêu hợp tác", "Article 1: Cooperation Objective", [
        ("双方共同的目标是，在平等互利的基础上，探索并推动中国科技企业，尤其是在教育科技、人工智能和环保科技等领域的优势企业，到越南进行市场拓展、技术合作与投资发展。乙方将作为专业的桥梁与顾问，协助这一过程，共同为越南经济和科技发展贡献力量。",
         "Mục tiêu chung của các Bên là, trên cơ sở bình đẳng và cùng có lợi, tìm hiểu và thúc đẩy các doanh nghiệp công nghệ Trung Quốc, đặc biệt là các doanh nghiệp có thế mạnh trong công nghệ giáo dục, trí tuệ nhân tạo và công nghệ môi trường, mở rộng thị trường, hợp tác công nghệ và đầu tư tại Việt Nam. Bên B sẽ đóng vai trò cầu nối và tư vấn chuyên nghiệp để hỗ trợ quá trình này, cùng đóng góp cho sự phát triển kinh tế và khoa học công nghệ của Việt Nam.",
         "The Parties’ common objective is, on the basis of equality and mutual benefit, to explore and promote market expansion, technology cooperation and investment in Vietnam by Chinese technology enterprises, particularly leading enterprises in education technology, artificial intelligence and environmental technology. Party B will serve as a professional bridge and advisor in this process, contributing jointly to Vietnam’s economic, scientific and technological development.")
    ]),
    ("第二条：合作范围", "Điều 2: Phạm vi hợp tác", "Article 2: Scope of Cooperation", [
        ("为实现上述目标，双方初步同意在以下领域探讨合作：", "Để đạt được mục tiêu nêu trên, các Bên bước đầu thống nhất tìm hiểu hợp tác trong các lĩnh vực sau:", "To achieve the above objective, the Parties preliminarily agree to explore cooperation in the following areas:"),
        ("1. 行业信息与政策共享：甲方可适时向乙方分享越南重点发展领域（如人工智能、环保、教育信息化等）的宏观政策方向、产业需求和市场准入规则。乙方可系统性地向甲方介绍中国在上述领域的成熟技术、成功应用案例和有出海意向的优质企业名录。",
         "1. Chia sẻ thông tin ngành và chính sách: Bên A có thể, vào thời điểm phù hợp, chia sẻ với Bên B định hướng chính sách vĩ mô, nhu cầu ngành và quy định tiếp cận thị trường của Việt Nam trong các lĩnh vực phát triển trọng điểm như trí tuệ nhân tạo, môi trường và số hóa giáo dục. Bên B có thể giới thiệu có hệ thống cho Bên A các công nghệ trưởng thành, trường hợp ứng dụng thành công và danh sách doanh nghiệp chất lượng cao của Trung Quốc có ý định mở rộng ra nước ngoài.",
         "1. Industry information and policy sharing: Party A may, as appropriate, share with Party B Vietnam’s macro policy directions, industry needs and market access rules in priority development areas such as artificial intelligence, environmental protection and education digitalization. Party B may systematically introduce to Party A mature Chinese technologies, successful application cases and lists of high-quality enterprises interested in overseas expansion."),
        ("2. 企业落地服务协同：乙方作为中国企业的咨询顾问，为其提供包括但不限于越南法律合规、知识产权保护框架搭建、跨境投融资架构设计、本地化运营支持等专业服务。甲方可在其职能范围内，为乙方推荐的中国企业提供有关越南投资环境、科研合作机会和潜在本地合作伙伴的参考信息与对接便利。",
         "2. Phối hợp dịch vụ hỗ trợ doanh nghiệp triển khai tại Việt Nam: Với vai trò tư vấn cho doanh nghiệp Trung Quốc, Bên B cung cấp các dịch vụ chuyên nghiệp bao gồm nhưng không giới hạn ở tuân thủ pháp luật Việt Nam, xây dựng khuôn khổ bảo hộ sở hữu trí tuệ, thiết kế cấu trúc đầu tư và tài chính xuyên biên giới, và hỗ trợ vận hành bản địa hóa. Trong phạm vi chức năng của mình, Bên A có thể cung cấp thông tin tham khảo và hỗ trợ kết nối về môi trường đầu tư, cơ hội hợp tác nghiên cứu và đối tác địa phương tiềm năng cho các doanh nghiệp Trung Quốc do Bên B giới thiệu.",
         "2. Coordination of enterprise market-entry services: As consultant to Chinese enterprises, Party B will provide professional services including, without limitation, Vietnamese legal compliance, intellectual property protection frameworks, cross-border investment and financing structures, and localized operational support. Within its functions, Party A may provide reference information and facilitate connections concerning Vietnam’s investment environment, research cooperation opportunities and potential local partners for Chinese enterprises recommended by Party B."),
        ("3. 共同举办活动：双方可探讨共同组织与“中越科技合作”相关的线上或线下论坛、企业对接会、路演和商务考察团。活动的具体主题、时间、地点和费用承担方式，将在后续单独商议。",
         "3. Đồng tổ chức sự kiện: Các Bên có thể tìm hiểu việc đồng tổ chức các diễn đàn trực tuyến hoặc trực tiếp, hội nghị kết nối doanh nghiệp, chương trình giới thiệu dự án và đoàn khảo sát kinh doanh liên quan đến “hợp tác khoa học công nghệ Việt Nam – Trung Quốc”. Chủ đề, thời gian, địa điểm và phương thức chịu chi phí cụ thể sẽ được thỏa thuận riêng sau đó.",
         "3. Joint events: The Parties may explore jointly organizing online or offline forums, business-matching meetings, roadshows and business study missions related to China–Vietnam technology cooperation. Specific themes, timing, venues and cost-sharing arrangements will be separately discussed."),
        ("4. 能力建设与人才交流：双方可探讨在人工智能、大数据、绿色技术等领域，共同设计短期培训或交流项目，以提升本地人才的专业技能。",
         "4. Nâng cao năng lực và trao đổi nhân tài: Các Bên có thể tìm hiểu việc cùng thiết kế các chương trình đào tạo hoặc trao đổi ngắn hạn trong các lĩnh vực như trí tuệ nhân tạo, dữ liệu lớn và công nghệ xanh nhằm nâng cao kỹ năng chuyên môn của nhân lực địa phương.",
         "4. Capacity building and talent exchange: The Parties may explore jointly designing short-term training or exchange programs in fields such as artificial intelligence, big data and green technology to enhance the professional skills of local talent.")
    ]),
    ("第三条：双方的角色与责任", "Điều 3: Vai trò và trách nhiệm", "Article 3: Roles and Responsibilities", [
        ("1. 甲方的角色：主要作为合作生态的本地网络支持方和战略方向的指引者。凭借其在越南官方和产学研界的公信力与影响力，为成功合作的开展和乙方的工作提供基础环境层面的支持。",
         "1. Vai trò của Bên A: Chủ yếu đóng vai trò hỗ trợ mạng lưới địa phương và định hướng chiến lược cho hệ sinh thái hợp tác. Dựa trên uy tín và ảnh hưởng trong các cơ quan, giới học thuật, nghiên cứu và doanh nghiệp tại Việt Nam, Bên A hỗ trợ tạo lập môi trường nền tảng cho việc triển khai hợp tác thành công và công việc của Bên B.",
         "1. Role of Party A: Primarily to serve as the local network supporter and strategic guide for the cooperation ecosystem. Leveraging its credibility and influence among Vietnamese institutions, academia, research and industry, Party A will support the foundational environment for successful cooperation and Party B’s work."),
        ("2. 乙方的角色：作为合作的具体执行和商业运营中枢，负责甄选、审核并推荐符合越南发展需求且具备实力的中国科技企业，并为中国企业出海全过程提供一站式、高质量咨询服务，管理项目商业风险，确保合作项目的专业性。",
         "2. Vai trò của Bên B: Là đầu mối triển khai cụ thể và vận hành thương mại của hợp tác; chịu trách nhiệm lựa chọn, thẩm định và giới thiệu các doanh nghiệp công nghệ Trung Quốc có năng lực, phù hợp với nhu cầu phát triển của Việt Nam; đồng thời cung cấp dịch vụ tư vấn một cửa, chất lượng cao trong toàn bộ quá trình doanh nghiệp Trung Quốc mở rộng ra nước ngoài, quản lý rủi ro thương mại và bảo đảm tính chuyên nghiệp của các dự án hợp tác.",
         "2. Role of Party B: To act as the operational and commercial hub of the cooperation; to select, review and recommend capable Chinese technology enterprises aligned with Vietnam’s development needs; and to provide one-stop, high-quality advisory services throughout their overseas expansion, manage commercial project risks and ensure the professionalism of cooperation projects.")
    ]),
    ("第四条：谅解与法律效力", "Điều 4: Tính chất và hiệu lực pháp lý", "Article 4: Understanding and Legal Effect", [
        ("本备忘录构成一份非正式、非官方的合作意向表达，而非具有法律约束力的合同。",
         "Biên bản ghi nhớ này là sự thể hiện ý định hợp tác mang tính phi chính thức, không phải văn kiện chính thức và không phải là hợp đồng có tính ràng buộc pháp lý.",
         "This Memorandum constitutes an informal and unofficial expression of cooperation intent and is not a legally binding contract."),
        ("1. 无约束力声明：本备忘录中的任何内容均不应被解释为在任何一方之间建立法律上的权利或义务关系。具体合作细节、责任归属、资金安排等，均需由相关方另行签署具有法律约束力的正式协议来确立。",
         "1. Tuyên bố không ràng buộc: Không nội dung nào trong Biên bản ghi nhớ này được hiểu là thiết lập quyền hoặc nghĩa vụ pháp lý giữa các Bên. Chi tiết hợp tác cụ thể, phân định trách nhiệm và thu xếp tài chính phải được xác lập bằng các thỏa thuận chính thức có tính ràng buộc pháp lý do các bên liên quan ký riêng.",
         "1. Non-binding statement: Nothing in this Memorandum shall be construed as creating legal rights or obligations between the Parties. Specific cooperation details, allocation of responsibilities and financial arrangements must be established through separate legally binding formal agreements signed by the relevant parties."),
        ("2. 非排他性：本备忘录不构成排他性合作安排，任何一方均有权与任何第三方开展相同或类似的合作。",
         "2. Không độc quyền: Biên bản ghi nhớ này không tạo thành thỏa thuận hợp tác độc quyền; mỗi Bên có quyền tiến hành hợp tác tương tự với bất kỳ bên thứ ba nào.",
         "2. Non-exclusivity: This Memorandum does not constitute an exclusive cooperation arrangement, and either Party may undertake the same or similar cooperation with any third party."),
        ("3. 无财务承诺：除非就特定项目另有书面约定，否则本备忘录项下的任何活动均不对任何一方构成直接拨款或其他财务承诺。双方原则上各自承担因参与本备忘录项下一般性活动而产生的费用。",
         "3. Không cam kết tài chính: Trừ khi có thỏa thuận khác bằng văn bản đối với một dự án cụ thể, không hoạt động nào theo Biên bản ghi nhớ này tạo thành nghĩa vụ cấp vốn trực tiếp hoặc cam kết tài chính khác đối với bất kỳ Bên nào. Về nguyên tắc, mỗi Bên tự chịu các chi phí phát sinh từ việc tham gia các hoạt động chung theo Biên bản ghi nhớ này.",
         "3. No financial commitment: Unless otherwise agreed in writing for a specific project, no activity under this Memorandum creates any direct funding obligation or other financial commitment for either Party. In principle, each Party will bear its own costs arising from participation in general activities under this Memorandum.")
    ]),
    ("第五条：合作成果的宣传与背书", "Điều 5: Truyền thông về kết quả hợp tác", "Article 5: Publicity of Cooperation Outcomes", [
        ("为推广合作并吸引更多优质企业参与，在不违反保密义务的前提下：",
         "Nhằm quảng bá hợp tác và thu hút thêm các doanh nghiệp chất lượng cao tham gia, với điều kiện không vi phạm nghĩa vụ bảo mật:",
         "To promote cooperation and attract more high-quality enterprises, subject to confidentiality obligations:"),
        ("1. 乙方可在其公司网站、宣传材料和商务洽谈中，如实陈述其与甲方存在本备忘录所描述的谅解合作关系，并可展示本备忘录的签署照片等公开信息，以此作为其信誉与能力的证明。",
         "1. Bên B có thể nêu trung thực trên trang web công ty, tài liệu quảng bá và trong đàm phán kinh doanh về mối quan hệ hợp tác theo Biên bản ghi nhớ này với Bên A, và có thể công bố hình ảnh ký kết cùng các thông tin công khai khác làm minh chứng cho uy tín và năng lực của mình.",
         "1. Party B may truthfully state on its corporate website, in promotional materials and in business discussions that it has the cooperation relationship with Party A described in this Memorandum, and may display signing photographs and other public information as evidence of its credibility and capabilities."),
        ("2. 甲方如需对外宣传双方的合作关系，应提前知会乙方。",
         "2. Nếu Bên A có nhu cầu truyền thông công khai về quan hệ hợp tác giữa các Bên, Bên A sẽ thông báo trước cho Bên B.",
         "2. If Party A wishes to publicize the cooperation relationship between the Parties, it shall notify Party B in advance.")
    ]),
    ("第六条：保密", "Điều 6: Bảo mật", "Article 6: Confidentiality", [
        ("在合作期间，一方（接收方）对于从另一方（披露方）获取的任何明确标识为“保密”的商业、技术或战略信息，应予以严格保密。未经披露方书面同意，不得向任何第三方泄露。此保密义务不因本备忘录的终止而失效。",
         "Trong thời gian hợp tác, một Bên (“Bên nhận”) phải bảo mật nghiêm ngặt mọi thông tin thương mại, kỹ thuật hoặc chiến lược được Bên kia (“Bên tiết lộ”) cung cấp và được đánh dấu rõ ràng là “bảo mật”. Bên nhận không được tiết lộ cho bất kỳ bên thứ ba nào nếu không có sự đồng ý bằng văn bản của Bên tiết lộ. Nghĩa vụ bảo mật này vẫn có hiệu lực sau khi Biên bản ghi nhớ chấm dứt.",
         "During the cooperation, a Party receiving any commercial, technical or strategic information clearly marked “Confidential” from the other Party shall keep such information strictly confidential and shall not disclose it to any third party without the disclosing Party’s written consent. This confidentiality obligation survives termination of this Memorandum.")
    ]),
    ("第七条：有效期与终止", "Điều 7: Thời hạn và chấm dứt", "Article 7: Term and Termination", [
        ("1. 本备忘录自双方授权代表签署之日起生效，有效期为三年。",
         "1. Biên bản ghi nhớ này có hiệu lực kể từ ngày được đại diện có thẩm quyền của cả hai Bên ký và có thời hạn ba năm.",
         "1. This Memorandum takes effect on the date it is signed by the authorized representatives of both Parties and remains effective for three years."),
        ("2. 有效期届满前，经双方协商一致，可书面续签。",
         "2. Trước khi hết thời hạn, Biên bản ghi nhớ có thể được gia hạn bằng văn bản theo thỏa thuận của các Bên.",
         "2. Before expiry, this Memorandum may be renewed in writing by mutual agreement of the Parties."),
        ("3. 任何一方均可提前三十日书面通知对方终止本备忘录。终止不影响已在进行中的具体项目（如有），也不影响第六条（保密）的效力。",
         "3. Mỗi Bên có thể chấm dứt Biên bản ghi nhớ này bằng cách thông báo bằng văn bản cho Bên kia trước ba mươi ngày. Việc chấm dứt không ảnh hưởng đến các dự án cụ thể đang triển khai (nếu có) hoặc hiệu lực của Điều 6 (Bảo mật).",
         "3. Either Party may terminate this Memorandum by giving the other Party thirty days’ prior written notice. Termination does not affect any ongoing specific project, if any, or the effectiveness of Article 6 (Confidentiality).")
    ]),
    ("第八条：签署、修改与其他", "Điều 8: Ký kết, sửa đổi và các vấn đề khác", "Article 8: Execution, Amendment and Miscellaneous", [
        ("1. 本备忘录以中文、越南文和英文起草，一式两份，双方各执一份。如对文本理解有歧义，以英文版本为准。",
         "1. Biên bản ghi nhớ này được lập bằng tiếng Trung, tiếng Việt và tiếng Anh thành hai bản, mỗi Bên giữ một bản. Trong trường hợp có khác biệt trong cách hiểu giữa các văn bản, bản tiếng Anh được ưu tiên áp dụng.",
         "1. This Memorandum is prepared in Chinese, Vietnamese and English in two originals, with each Party retaining one original. In the event of any discrepancy in interpretation, the English version shall prevail."),
        ("2. 对本备忘录的任何修改或补充，须经双方书面同意并签署补充文件后方为有效。",
         "2. Mọi sửa đổi hoặc bổ sung đối với Biên bản ghi nhớ này chỉ có hiệu lực khi được các Bên thống nhất bằng văn bản và ký văn bản bổ sung.",
         "2. Any amendment or supplement to this Memorandum is effective only when agreed in writing by the Parties and executed in a supplementary document."),
        ("3. 因本备忘录或其履行产生的或与之相关的任何争议，双方应首先通过友好协商解决。",
         "3. Mọi tranh chấp phát sinh từ hoặc liên quan đến Biên bản ghi nhớ này hoặc việc thực hiện Biên bản ghi nhớ trước hết sẽ được các Bên giải quyết thông qua tham vấn hữu nghị.",
         "3. Any dispute arising out of or relating to this Memorandum or its performance shall first be resolved by the Parties through friendly consultation."),
        ("兹证明，双方已促使其各自正式授权代表于文首所载日期签署本谅解备忘录。",
         "Để làm bằng chứng, các Bên đã yêu cầu đại diện được ủy quyền hợp lệ của mình ký Biên bản ghi nhớ này vào ngày nêu tại phần đầu.",
         "In witness whereof, the Parties have caused their duly authorized representatives to sign this Memorandum on the date first written above.")
    ]),
]


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.27)
sec.page_height = Inches(11.69)
sec.top_margin = Inches(0.65)
sec.bottom_margin = Inches(0.65)
sec.left_margin = Inches(0.65)
sec.right_margin = Inches(0.65)
sec.header_distance = Inches(0.3)
sec.footer_distance = Inches(0.3)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = header.add_run("中越英三语打印版  |  Bản in ba ngôn ngữ  |  Trilingual Print Version")
set_font(r, 8.5, color=GRAY)
add_page_number(sec.footer.paragraphs[0])

add_title(doc)
add_intro_table(doc)

for zh, vi, en, rows in sections:
    add_heading(doc, zh, vi, en)
    add_trilingual_table(doc, rows)

doc.add_page_break()
add_heading(doc, "签署页", "Trang ký kết", "Signature Page")
add_trilingual_table(doc, [(
    "双方授权代表在下方签署。本签署页为本备忘录不可分割的一部分。",
    "Đại diện được ủy quyền của các Bên ký tên dưới đây. Trang ký kết này là một phần không tách rời của Biên bản ghi nhớ.",
    "The authorized representatives of the Parties sign below. This signature page forms an integral part of this Memorandum."
)])

sig = doc.add_table(rows=1, cols=2)
sig.style = "Table Grid"
texts = [
    "代表甲方 / ĐẠI DIỆN BÊN A / FOR PARTY A\n\n\n签字 / Chữ ký / Signature: __________________________\n\n姓名 / Họ tên / Name: ______________________________\n\n职务 / Chức vụ / Title: _____________________________\n\n日期 / Ngày / Date: ________________________________",
    "代表乙方 / ĐẠI DIỆN BÊN B / FOR PARTY B\n新瑞拓海有限公司\nShinered Oversea Co., Ltd.\n\n\n签字 / Chữ ký / Signature: __________________________\n\n姓名 / Họ tên / Name: ______________________________\n\n职务 / Chức vụ / Title: _____________________________\n\n日期 / Ngày / Date: ________________________________",
]
for i, text in enumerate(texts):
    p = sig.rows[0].cells[i].paragraphs[0]
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, 10, i == 1)
table_geometry(sig, [4680, 4680])

doc.core_properties.title = "中越合作备忘录（中越英三语打印版）"
doc.core_properties.subject = "Memorandum of Understanding - Chinese, Vietnamese and English"
doc.core_properties.author = "Shinered Oversea Co., Ltd."
doc.save(OUT)
print(OUT)
