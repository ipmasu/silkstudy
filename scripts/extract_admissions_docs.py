from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

import pdfplumber
from docx import Document


SOURCE_DIR = Path(r"C:\Users\马肃\Downloads\招生简章\已经整理上传")
OUTPUT_DIR = Path("outputs/admissions-extract")
MAX_PDF_PAGES = 12


KEYWORDS = [
    "学费",
    "申请费",
    "报名费",
    "住宿费",
    "保险",
    "HSK",
    "汉语",
    "中文授课",
    "本科",
    "申请时间",
    "截止",
    "奖学金",
    "CSCA",
    "护照",
    "高中",
    "非中国籍",
    "contact",
    "tuition",
    "application fee",
    "deadline",
    "scholarship",
    "Chinese-taught",
    "undergraduate",
]


def normalize(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(path: Path) -> str:
    chunks: list[str] = []
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages[:MAX_PDF_PAGES], start=1):
            text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            if text.strip():
                chunks.append(f"\n--- page {index} ---\n{text}")
    return normalize("\n".join(chunks))


def extract_docx(path: Path) -> str:
    doc = Document(path)
    chunks: list[str] = []
    for para in doc.paragraphs:
        value = para.text.strip()
        if value:
            chunks.append(value)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    return normalize("\n".join(chunks))


def keyword_snippets(text: str) -> list[str]:
    snippets: list[str] = []
    seen: set[str] = set()
    lowered = text.lower()
    for keyword in KEYWORDS:
        start = 0
        needle = keyword.lower()
        while len(snippets) < 80:
            index = lowered.find(needle, start)
            if index == -1:
                break
            left = max(0, index - 180)
            right = min(len(text), index + 420)
            snippet = normalize(text[left:right])
            if snippet not in seen:
                seen.add(snippet)
                snippets.append(snippet)
            start = index + len(keyword)
    return snippets


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    records = []
    for path in sorted(SOURCE_DIR.iterdir(), key=lambda item: item.name):
        if path.suffix.lower() not in {".pdf", ".docx"}:
            continue
        if path.name == "已经整理上传.rar":
            continue
        try:
            text = extract_pdf(path) if path.suffix.lower() == ".pdf" else extract_docx(path)
            text_path = OUTPUT_DIR / f"{path.stem}.txt"
            text_path.write_text(text, encoding="utf-8")
            records.append(
                {
                    "file": str(path),
                    "textFile": str(text_path),
                    "chars": len(text),
                    "snippets": keyword_snippets(text),
                }
            )
            summary_path = OUTPUT_DIR / "summary.json"
            summary_path.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception as exc:
            records.append({"file": str(path), "error": repr(exc)})
            summary_path = OUTPUT_DIR / "summary.json"
            summary_path.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")

    summary_path = OUTPUT_DIR / "summary.json"
    summary_path.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {summary_path} with {len(records)} records")


if __name__ == "__main__":
    main()
